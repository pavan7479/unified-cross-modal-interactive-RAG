import io
import requests
import pandas as pd
from typing import List
import time
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup
from langchain_community.document_loaders import YoutubeLoader
from langchain_core.documents import Document

from config import Config
from utils.logger import get_logger

logger = get_logger(__name__)


# ==========================================================
# GENERIC FILE LOADER (PDF, TXT, DOCX, CSV)
# ==========================================================

def load_uploaded_file(uploaded_file) -> List[Document]:
    """
    Handles PDF, TXT, DOCX, CSV from Streamlit uploader.
    All processed in memory.
    """

    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return load_pdf(uploaded_file)

    elif file_name.endswith(".txt"):
        return load_txt(uploaded_file)

    elif file_name.endswith(".docx"):
        return load_docx(uploaded_file)

    elif file_name.endswith(".csv"):
        return load_csv(uploaded_file)

    else:
        raise ValueError(f"Unsupported file type: {file_name}")


# ==========================================================
# PDF
# ==========================================================

def load_pdf(uploaded_file) -> List[Document]:
    import fitz  # PyMuPDF

    docs = []
    pdf_bytes = uploaded_file.read()
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page_num, page in enumerate(pdf, start=1):
        text = page.get_text("text")
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source_type": "pdf",
                        "source": uploaded_file.name,
                        "page": page_num,
                    },
                )
            )

    logger.info(f"Loaded {len(docs)} pages from PDF")
    return docs


# ==========================================================
# TXT
# ==========================================================

def load_txt(uploaded_file) -> List[Document]:
    content = uploaded_file.read().decode("utf-8", errors="ignore")

    return [
        Document(
            page_content=content,
            metadata={
                "source_type": "txt",
                "source": uploaded_file.name,
            },
        )
    ]


# ==========================================================
# DOCX
# ==========================================================

def load_docx(uploaded_file) -> List[Document]:
    from docx import Document as DocxDocument

    file_stream = io.BytesIO(uploaded_file.read())
    doc = DocxDocument(file_stream)

    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    return [
        Document(
            page_content=text,
            metadata={
                "source_type": "docx",
                "source": uploaded_file.name,
            },
        )
    ]


# ==========================================================
# CSV
# ==========================================================

def load_csv(uploaded_file) -> List[Document]:
    df = pd.read_csv(uploaded_file)

    text = df.to_string(index=False)

    return [
        Document(
            page_content=text,
            metadata={
                "source_type": "csv",
                "source": uploaded_file.name,
            },
        )
    ]


# ==========================================================
# YOUTUBE (Transcript Only)
# ==========================================================

# ==========================================================
# YOUTUBE (LangChain Loader)
# ==========================================================

from langchain_community.document_loaders import YoutubeLoader


def load_youtube(url: str) -> List[Document]:
    """
    Loads YouTube transcript using LangChain's YoutubeLoader.
    Uses transcript only (no audio processing).
    """

    try:
        loader = YoutubeLoader.from_youtube_url(
            url,
            add_video_info=False,
            language=["en"],
        )

        documents = loader.load()

        # Enrich metadata
        for doc in documents:
            doc.metadata["source_type"] = "youtube"

        logger.info(f"Loaded YouTube transcript: {url}")

        return documents

    except Exception as e:
        logger.warning(f"YouTube loading failed: {e}")
        raise ValueError(
            "Could not fetch YouTube transcript. "
            "YouTube may be blocking cloud requests."
        )

# ==========================================================
# SIMPLE WEB LOADER (NO HEAVY RECURSION)
# ==========================================================




def load_web(url: str, depth: int = 0, visited=None) -> List[Document]:
    if visited is None:
        visited = set()

    if depth > Config.MAX_CRAWL_DEPTH:
        return []

    if url in visited:
        return []

    visited.add(url)

    try:
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "en-US,en;q=0.9",
        }

        response = requests.get(
            url,
            headers=headers,
            timeout=Config.REQUEST_TIMEOUT,
        )

        response.raise_for_status()

        soup = BeautifulSoup(response.text, "html.parser")

        # Remove noisy tags
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        text = soup.get_text(separator="\n")
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())

        documents = [
            Document(
                page_content=text,
                metadata={
                    "source_type": "web",
                    "source": url,
                    "depth": depth,
                },
            )
        ]

        # ----------------------------
        # Controlled Recursive Crawling
        # ----------------------------
        if depth < Config.MAX_CRAWL_DEPTH:

            base_domain = urlparse(url).netloc
            links = []

            for a in soup.find_all("a", href=True):
                absolute_url = urljoin(url, a["href"])

                # Same domain only
                if urlparse(absolute_url).netloc == base_domain:
                    if absolute_url not in visited:
                        links.append(absolute_url)

            # Limit crawl size
            for link in links[: Config.MAX_CRAWL_PAGES]:
                time.sleep(0.5)  # polite delay
                documents.extend(load_web(link, depth + 1, visited))

        return documents

    except requests.exceptions.HTTPError as e:
        logger.warning(f"HTTP error while scraping {url}: {e}")
        return []

    except Exception as e:
        logger.warning(f"Web scrape failed for {url}: {e}")
        return []
